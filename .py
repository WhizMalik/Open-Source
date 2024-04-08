import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from heapq import nlargest
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from PyPDF2 import PdfReader

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def generate_notes(transcript, num_sentences=5):
    # Tokenize the transcript into sentences
    sentences = sent_tokenize(transcript)

    # Tokenize each sentence into words and remove stopwords
    stop_words = set(stopwords.words('english'))
    word_frequencies = {}
    for sentence in sentences:
        words = word_tokenize(sentence.lower())
        for word in words:
            if word not in stop_words and word.isalnum():
                if word not in word_frequencies:
                    word_frequencies[word] = 1
                else:
                    word_frequencies[word] += 1

    # Calculate sentence scores based on word frequencies
    sentence_scores = {}
    for i, sentence in enumerate(sentences):
        for word in word_tokenize(sentence.lower()):
            if word in word_frequencies:
                if i not in sentence_scores:
                    sentence_scores[i] = word_frequencies[word]
                else:
                    sentence_scores[i] += word_frequencies[word]

    # Get the top N sentences with the highest scores
    summary_sentences = nlargest(num_sentences, sentence_scores, key=sentence_scores.get)

    # Sort the summary sentences in the order they appear in the original transcript
    summary_sentences.sort()

    # Join the summary sentences to form the notes
    notes = " ".join([sentences[i] for i in summary_sentences])

    return notes

def summarize(text, num_sentences=3):
    # Tokenize the text into sentences
    sentences = sent_tokenize(text)

    # Calculate sentence scores based on word frequencies
    word_frequencies = {}
    for sentence in sentences:
        words = word_tokenize(sentence.lower())
        for word in words:
            if word not in stopwords.words('english') and word.isalnum():
                if word not in word_frequencies:
                    word_frequencies[word] = 1
                else:
                    word_frequencies[word] += 1

    sentence_scores = {}
    for i, sentence in enumerate(sentences):
        for word in word_tokenize(sentence.lower()):
            if word in word_frequencies:
                if i not in sentence_scores:
                    sentence_scores[i] = word_frequencies[word]
                else:
                    sentence_scores[i] += word_frequencies[word]

    # Get the top N sentences with the highest scores
    summary_sentences = nlargest(num_sentences, sentence_scores, key=sentence_scores.get)

    # Sort the summary sentences in the order they appear in the original text
    summary_sentences.sort()

    # Join the summary sentences to form the summary
    summary = " ".join([sentences[i] for i in summary_sentences])

    return summary

# Example usage
pdf_path = "CCN M4_up.pdf"  # Replace with the path to your PDF file

# Extract text from the PDF
transcript = extract_text_from_pdf(pdf_path)

notes = generate_notes(transcript)
summary = summarize(transcript)

# Create a new Word document
document = Document()

# Add a heading for the notes
notes_heading = document.add_heading("Important Points", level=1)
notes_heading.style.font.color.rgb = RGBColor(0, 0, 255)  # Set color to blue

# Add the notes to the document
for sentence in sent_tokenize(notes):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(sentence)
    run.bold = True
    run.font.size = Pt(12)

# Add a heading for the summary
summary_heading = document.add_heading("Summary", level=1)
summary_heading.style.font.color.rgb = RGBColor(255, 0, 0)  # Set color to red

# Add the summary to the document
summary_paragraph = document.add_paragraph(summary)
summary_paragraph.style.font.size = Pt(12)

# Save the document
document.save("gen_notes.docx")