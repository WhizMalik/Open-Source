import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from heapq import nlargest
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from PyPDF2 import PdfReader

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def preprocess_text(text):
    # Remove non-alphanumeric characters and convert to lowercase
    text = re.sub(r'\W+', ' ', text.lower())
    return text

def generate_exam_notes(transcript, num_sentences):
    # Tokenize the transcript into sentences
    sentences = sent_tokenize(transcript)
    print("Number of sentences:", len(sentences))

    # Tokenize each sentence into words and remove stopwords
    stop_words = set(stopwords.words('english'))
    word_frequencies = {}
    for sentence in sentences:
        words = word_tokenize(sentence)
        for word in words:
            if word not in stop_words and word.isalnum():
                if word not in word_frequencies:
                    word_frequencies[word] = 1
                else:
                    word_frequencies[word] += 1

    # Calculate sentence scores based on word frequencies, sentence position, and exam keywords
    sentence_scores = {}
    exam_keywords = ["define", "explain", "describe", "analyze", "compare", "contrast", "evaluate"]
    for i, sentence in enumerate(sentences):
        for word in word_tokenize(sentence.lower()):
            if word in word_frequencies:
                if i not in sentence_scores:
                    sentence_scores[i] = word_frequencies[word]
                else:
                    sentence_scores[i] += word_frequencies[word]
        sentence_scores[i] /= (i + 1)  # Give higher scores to sentences appearing earlier
        
        # Increase score for sentences containing exam keywords
        for keyword in exam_keywords:
            if keyword in sentence.lower():
                sentence_scores[i] *= 1.5
                break

    print("Sentence scores:", sentence_scores)

    # Get the top N sentences with the highest scores
    summary_sentences = nlargest(num_sentences, sentence_scores, key=sentence_scores.get)
    print("Summary sentences:", summary_sentences)

    # Sort the summary sentences in the order they appear in the original transcript
    summary_sentences.sort()

    # Join the summary sentences to form the exam notes
    exam_notes = " ".join([sentences[i] for i in summary_sentences])
    print("Exam notes:", exam_notes)

    return exam_notes

# Example usage
pdf_path = "CCN M4_up.pdf"  # Replace with the path to your PDF file

# Extract text from the PDF
transcript = extract_text_from_pdf(pdf_path)
print("Extracted text:", transcript)

# Preprocess the text
preprocessed_text = preprocess_text(transcript)
print("Preprocessed text:", preprocessed_text)

# Calculate the number of sentences for exam notes based on the desired output size
total_sentences = len(sent_tokenize(preprocessed_text))
notes_ratio = 0.3
notes_sentences = int(total_sentences * notes_ratio)
print("Number of sentences for exam notes:", notes_sentences)

# Generate exam notes
exam_notes = generate_exam_notes(preprocessed_text, num_sentences=notes_sentences)

# Create a new Word document
document = Document()

# Add a heading for the exam notes
notes_heading = document.add_heading("Exam Notes", level=1)
notes_heading.style.font.color.rgb = RGBColor(0, 0, 255)  # Set color to blue

# Add the exam notes to the document
for sentence in sent_tokenize(exam_notes):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(sentence)
    run.bold = True
    run.font.size = Pt(12)

# Save the document
document.save("exam_notes.docx")